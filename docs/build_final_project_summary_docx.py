from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "FINAL_PROJECT_SUMMARY.docx"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = str(text)
        shade(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = str(text)
    doc.add_paragraph()


def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    for line in text.strip("\n").splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)


def add_bullets(doc, values):
    for value in values:
        doc.add_paragraph(value, style="List Bullet")


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("nanoGPT-ZYNQ 最终工程总结")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Microsoft YaHei"
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("75 MHz / INT8 / 6 层 Transformer / PS-PL 协同推理").italic = True
    doc.add_paragraph()

    doc.add_heading("1. 项目结论", level=1)
    doc.add_paragraph(
        "本工程完成了部署在 Zynq-7020 上的 6 层 INT8 nanoGPT 推理系统。最终板级版本采用 "
        "75 MHz 单时钟、DDR 存储权重和中间数据、PS 端串口交互、PL 端 Transformer 计算和 DDR K/V Cache。"
        "模型为 Shakespeare 字符级模型，输入英文字符并输出下一个英文字符。"
    )
    doc.add_paragraph(
        "最终版本已完成 200 字符串口生成和六层 hidden 对齐测试。交付包只包含最终版本，未包含旧双时钟工程、"
        "历史 .bak 文件或 Vivado 缓存。"
    )

    doc.add_heading("2. 最终配置", level=1)
    add_table(doc, ["项目", "最终值"], [
        ["FPGA", "XC7Z020 (Zynq-7020)"],
        ["PL 时钟", "75 MHz"],
        ["模型结构", "6 层 Transformer, d_model=384, context=256"],
        ["词表", "65 个 Shakespeare 字符级 token"],
        ["数值格式", "INT8 权重/激活，定点缩放参数位于 DDR"],
        ["生成策略", "Greedy argmax"],
        ["串口", "板载 FTDI UART, 115200 baud, 8N1"],
        ["最长生成", "200 字符"],
        ["最终加速", "FFN 跨组预取 + FFN 32 路并行 + DDR K/V Cache"],
    ])

    doc.add_heading("3. PS-PL 数据流", level=1)
    add_code(doc, """PC UART prompt
    -> PS: character encode / token id / embedding quantization
    -> DDR: input hidden, weights, scales, LUTs, K/V cache
    -> PL: 6-layer INT8 Transformer and LM-head argmax
    -> PS: token decode
    -> UART character output""")
    add_table(doc, ["DDR 数据", "用途"], [
        ["input.bin", "输入 hidden 和初始测试向量"],
        ["weights.bin", "INT8 权重，约 10.15 MB"],
        ["scales.bin", "各层定点缩放参数"],
        ["luts.bin", "softmax/GELU 查找表"],
        ["golden_final.bin", "最终 hidden golden 对照数据"],
        ["K/V Cache", "各层 K/V 缓存；后续 token 尽量只计算新增行"],
    ])

    doc.add_heading("4. 重要源码", level=1)
    add_table(doc, ["文件", "主要职责"], [
        ["fpga/rtl/hls_kernel_chain_axis_full_only_core.v", "PL 核心：AXI4 DDR、六层 INT8 状态机、FFN32 和 LM head。"],
        ["ps/main.c", "UART、65 字符 token 编解码、embedding、K/V Cache 调度与输出解码。"],
        ["ps/run_ps_mailbox_runner.tcl", "JTAG 下载 bitstream、DDR 镜像和 ELF，启动 UART。"],
        ["fpga/scripts/setup_vivado_project_ddr.tcl", "Vivado 工程、PS7、DDR、AXI 与 PL 核连接。"],
        ["python/tools/eval_int8_reference.py", "FP32/INT8 质量对比。"],
        ["python/tools/pack_int8_full_ddr_image.py", "打包最终 DDR 运行镜像。"],
    ])

    doc.add_paragraph("字符编码核心（ps/main.c）：")
    add_code(doc, """static int encode_char(char c)
{
    if (c == '\\n') return 0;
    if (c == ' ') return 1;
    if (c >= 'A' && c <= 'Z') return 13 + (c - 'A');
    if (c >= 'a' && c <= 'z') return 39 + (c - 'a');
    return -1;
}""")
    doc.add_paragraph("K/V Cache 和字符输出核心（ps/main.c）：")
    add_code(doc, """rc = run_full_model_range(n, row_start);
rc = pl_lm_head_argmax_row(LAYER_A_BASE + ((n - 1u) * D_MODEL), &tok);
mailbox_write(MAILBOX_TOKEN_WORD_BASE + generated, (uint32_t)tok);
mailbox_write(MAILBOX_CHAR_WORD_BASE + generated,
              (tok < VOCAB_SIZE) ? (uint32_t)g_itos[tok] : (uint32_t)'?');
tokens[n++] = tok;""")
    doc.add_paragraph("FFN32 并行结构（核心 RTL）：")
    add_code(doc, """logic signed [17:0] ffn_mul_b0;   // ... b1 to b30
logic signed [17:0] ffn_mul_b31;
(* use_dsp = \"yes\" *) logic signed [42:0] ffn_prod0;  // ... prod1 to prod30
(* use_dsp = \"yes\" *) logic signed [42:0] ffn_prod31;""")

    doc.add_heading("5. INT8 质量指标", level=1)
    doc.add_paragraph("固定 INT8 质量门限测试覆盖 8 个数据块、共 2048 个预测位置。")
    add_table(doc, ["指标", "FP32", "INT8 / 对比结果"], [
        ["Loss", "1.576245", "1.628690"],
        ["Perplexity (PPL)", "4.836758", "5.097195"],
        ["PPL 回退", "-", "5.385%"],
        ["Logits MAE", "-", "0.344568"],
        ["Top-1 match", "-", "1837 / 2048 = 89.697%"],
        ["门限判断", "-", "通过：PPL 回退小于 10%"],
    ])
    doc.add_paragraph(
        "全硬件语义质量记录中，FP32 PPL 为 264.060077，INT8 PPL 为 284.159996，"
        "PPL 回退为 7.612%，同样低于 10% 门限。"
    )

    doc.add_heading("6. 优化结果", level=1)
    add_table(doc, ["步骤", "实现", "200 字符耗时", "结论"], [
        ["1", "FFN 跨组预取", "99.080 s", "正确，逐字符一致"],
        ["2", "75/100 MHz 双时钟 DDR", "113.707 s", "正确但异步 FIFO 使短事务变慢"],
        ["3", "FFN 32 路并行", "71.075 s", "最终采用，较 FFN16 提升 1.39 倍"],
    ])
    doc.add_paragraph("最终 FFN 仿真周期从 493,872 降至 327,864，减少 33.6%。")

    doc.add_heading("7. 时序和资源", level=1)
    add_table(doc, ["资源", "使用量", "可用量", "占用率"], [
        ["Slice LUT", "25,678", "53,200", "48.27%"],
        ["Slice Register", "29,482", "106,400", "27.71%"],
        ["Block RAM Tile", "109", "140", "77.86%"],
        ["DSP", "102", "220", "46.36%"],
    ])
    add_table(doc, ["时序指标", "结果"], [
        ["WNS", "+0.332 ns"],
        ["TNS", "0.000 ns"],
        ["WHS", "+0.036 ns"],
        ["约束", "全部满足"],
    ])
    doc.add_paragraph("BRAM 是当前最紧张的资源，仍有 22.14% 余量；DSP 使用率为 46.36%。")

    doc.add_heading("8. 板级验收", level=1)
    doc.add_paragraph("串口命令：")
    add_code(doc, "200:everything with a man")
    doc.add_paragraph("板端 200 字符输出开头：")
    add_code(doc, """that we have stood
The seal of the sea of the war, the world begins
Of the seass of the seasons of the world,
Which the seals of the sea of the world,""")
    add_bullets(doc, [
        "串口日志总耗时：71.075 s。",
        "输出逐字符流式返回。",
        "六层计算结束后最后一行 hidden：96/96 个 32 位字一致。",
        "时序、资源、功能和串口交互均通过最终验收。",
    ])

    doc.add_heading("9. 交付和复现", level=1)
    add_bullets(doc, [
        "使用 fpga/overlay/system/system.bit 和 system.hwh 配置 PL。",
        "通过 ps/run_ps_mailbox_runner.tcl 下载 DDR 镜像并启动 ps_mailbox_runner.elf。",
        "UART 设置为 115200, 8N1，发送 200:everything with a man 后回车。",
        "对照 tests/ 和 artifacts/reports/ 中的最终测试、时序和资源报告。",
    ])

    doc.add_heading("10. 原始证据文件", level=1)
    add_bullets(doc, [
        "artifacts/reports/optimization_123_results_20260712.md",
        "artifacts/reports/timing_step3_ffn32_resetfix_true75_post_route.rpt",
        "artifacts/reports/utilization_step3_ffn32_resetfix_true75_post_route.rpt",
        "tests/step3_ffn32_resetfix_200.raw.txt",
        "metrics/int8_bittrue_fixed_eval_metrics.json",
        "metrics/hardware_semantics_quality_metrics.md",
    ])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
